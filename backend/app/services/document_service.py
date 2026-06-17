from docx import Document as DocxDocument

from langchain_core.documents import (
    Document
)


def extract_docx_text(
    file_path
):

    doc = DocxDocument(
        file_path
    )

    text = "\n".join(
        [
            paragraph.text
            for paragraph in doc.paragraphs
        ]
    )

    print("=" * 50)
    print("DOCX LENGTH:", len(text))
    print(text[:500])
    print("=" * 50)

    return [
        Document(
            page_content=text
        )
    ]


def extract_md_text(
    file_path
):

    with open(
        file_path,
        "r",
        encoding="utf-8"
    ) as file:

        text = file.read()

    print("=" * 50)
    print("MD LENGTH:", len(text))
    print(text[:500])
    print("=" * 50)

    return [
        Document(
            page_content=text
        )
    ]


def extract_txt_text(
    file_path
):

    with open(
        file_path,
        "r",
        encoding="utf-8"
    ) as file:

        text = file.read()

    print("=" * 50)
    print("TXT LENGTH:", len(text))
    print(text[:500])
    print("=" * 50)

    return [
        Document(
            page_content=text
        )
    ]